"""Extract text from DOCX (Word) for ontology discovery."""
from typing import Any, Dict, List, Tuple

from app.models.ontology import OntologyEvidence, SourceArtifact


class DocxProcessor:
    """Extract text and paragraph structure from DOCX."""

    def process(self, artifact: SourceArtifact) -> Tuple[str, List[Dict[str, Any]], List[OntologyEvidence]]:
        """
        Process a DOCX artifact. Returns:
        - full_text
        - sections: list of {heading, content, page_number} (page_number 0 for docx)
        - evidence_list
        """
        full_text_parts: List[str] = []
        sections: List[Dict[str, Any]] = []
        evidence_list: List[OntologyEvidence] = []

        try:
            import docx
            doc = docx.Document(artifact.file_path)
            for i, para in enumerate(doc.paragraphs):
                text = (para.text or "").strip()
                if not text:
                    continue
                full_text_parts.append(text)
                style = (para.style and para.style.name or "").lower()
                is_heading = "heading" in style or "title" in style
                sections.append({
                    "heading": text[:80] if is_heading else "Body",
                    "content": text,
                    "page_number": 0,
                })
                evidence_list.append(
                    OntologyEvidence(
                        id=f"evt_{artifact.id}_p{i}_{len(evidence_list)}",
                        artifact_id=artifact.id,
                        artifact_type="docx",
                        page_number=None,
                        text_snippet=text[:500],
                        extraction_stage="docx_processor",
                    )
                )
            if not full_text_parts and doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
                        if row_text:
                            full_text_parts.append(row_text)
                            sections.append({"heading": "Table", "content": row_text, "page_number": 0})
        except ImportError:
            full_text_parts.append("[python-docx not installed; install it to process Word files]")
        except Exception as e:
            full_text_parts.append(f"[DOCX read error: {e}]")

        full_text = "\n\n".join(full_text_parts)
        return full_text, sections, evidence_list
